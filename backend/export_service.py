"""
Export service: chuyển nội dung Markdown sang MD / PDF / DOCX.
"""

import io
import re
import os
import glob


# ═════════════════════════════════════════════════════════════
# 1) MARKDOWN
# ═════════════════════════════════════════════════════════════
def export_markdown(content_md: str) -> bytes:
    """Trả về bytes UTF-8 của nội dung markdown."""
    return content_md.encode("utf-8")


# ═════════════════════════════════════════════════════════════
# 2) PDF  — dùng ReportLab
# ═════════════════════════════════════════════════════════════
def _register_pdf_fonts():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    base_dir = os.path.dirname(os.path.abspath(__file__))
    candidates = [
        (os.path.join(base_dir, "fonts", "Inter-Regular.ttf"), "Inter"),
        (os.path.join(base_dir, "fonts", "Inter-Bold.ttf"),    "Inter-Bold"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",      "DejaVuSans"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVuSans-Bold"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",  "DejaVuSansMono"),
    ]

    registered = {}
    for path, name in candidates:
        if not glob.glob(path):
            continue
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            registered[name] = path
        except Exception:
            pass

    # Đăng ký DejaVuSansMono riêng cho code block
    mono_path = "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf"
    if os.path.exists(mono_path) and "DejaVuSansMono" not in registered:
        try:
            pdfmetrics.registerFont(TTFont("DejaVuSansMono", mono_path))
        except Exception:
            pass

    # Ưu tiên Inter → DejaVuSans → fallback Helvetica
    for base, bold in [
        ("Inter",      "Inter-Bold"),
        ("DejaVuSans", "DejaVuSans-Bold"),
    ]:
        if base in registered:
            try:
                from reportlab.pdfbase.pdfmetrics import registerFontFamily
                bold_name = bold if bold in registered else base
                registerFontFamily(base, normal=base, bold=bold_name,
                                   italic=base, boldItalic=bold_name)
            except Exception:
                pass
            return base, bold if bold in registered else base

    return "Helvetica", "Helvetica-Bold"


def export_pdf(title: str, content_md: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Preformatted,
            ListFlowable, ListItem, HRFlowable,
        )
        from reportlab.lib.enums import TA_LEFT
    except ImportError:
        raise RuntimeError("Thư viện reportlab chưa được cài đặt.")

    FONT_NORMAL, FONT_BOLD = _register_pdf_fonts()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
    )

    styles = getSampleStyleSheet()
    PRIMARY = colors.HexColor("#4F46E5")

    style_title = ParagraphStyle(
        "DocTitle", parent=styles["Title"],
        textColor=PRIMARY, fontSize=20,
        spaceAfter=16, fontName=FONT_BOLD,
    )
    style_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        textColor=PRIMARY, fontSize=16,
        spaceBefore=12, spaceAfter=6, fontName=FONT_BOLD,
    )
    style_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        textColor=PRIMARY, fontSize=14,
        spaceBefore=10, spaceAfter=5, fontName=FONT_BOLD,
    )
    style_h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontSize=12, spaceBefore=8, spaceAfter=4, fontName=FONT_BOLD,
    )
    style_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName=FONT_NORMAL, fontSize=11, leading=16, spaceAfter=6,
    )
    style_code = ParagraphStyle(
        "Code", parent=styles["Code"],
        fontName="DejaVuSansMono", fontSize=9,
        backColor=colors.HexColor("#F3F4F6"),
        borderPadding=8, leading=13, spaceAfter=8,
    )
    style_bullet = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontName=FONT_NORMAL, fontSize=11,
        leading=16, leftIndent=16, spaceAfter=4,
    )

    story = []
    story.append(Paragraph(_escape_html(title), style_title))
    story.append(HRFlowable(color=PRIMARY, thickness=1.5, spaceAfter=12))

    lines = content_md.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_buf), style_code))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            lvl = len(m.group(1))
            text = _escape_html(m.group(2).strip())
            if lvl == 1:
                story.append(Paragraph(text, style_h1))
            elif lvl == 2:
                story.append(Paragraph(text, style_h2))
            else:
                story.append(Paragraph(text, style_h3))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            story.append(HRFlowable(
                color=colors.HexColor("#E5E7EB"), thickness=1, spaceAfter=8
            ))
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                text = re.sub(r"^\s*[-*+]\s+", "", lines[i])
                items.append(ListItem(Paragraph(_inline_html(text), style_bullet)))
                i += 1
            story.append(ListFlowable(items, bulletType="bullet", leftIndent=16))
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                text = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                items.append(ListItem(Paragraph(_inline_html(text), style_bullet)))
                i += 1
            story.append(ListFlowable(items, bulletType="1", leftIndent=16))
            continue

        # Blockquote
        if line.startswith(">"):
            text = re.sub(r"^>\s?", "", line)
            bq_style = ParagraphStyle(
                "BQ", parent=style_body,
                leftIndent=20, textColor=colors.HexColor("#6B7280"),
            )
            story.append(Paragraph(_inline_html(text), bq_style))
            i += 1
            continue

        # Empty line
        if not line.strip():
            story.append(Spacer(1, 6))
            i += 1
            continue

        # Normal paragraph
        story.append(Paragraph(_inline_html(line), style_body))
        i += 1

    # Close open code block
    if in_code and code_buf:
        story.append(Preformatted("\n".join(code_buf), style_code))

    doc.build(story)
    return buf.getvalue()


def _escape_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_html(text: str) -> str:
    """Convert **bold**, *italic*, `code` sang ReportLab HTML tags."""
    text = _escape_html(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"`(.+?)`", r'<font name="DejaVuSansMono" size="9">\1</font>', text)
    return text


# ═════════════════════════════════════════════════════════════
# 3) DOCX — python-docx
# ═════════════════════════════════════════════════════════════
def export_docx(title: str, content_md: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("Thư viện python-docx chưa được cài đặt.")

    doc = Document()
    _set_doc_default_font(doc, "DejaVuSans", 11)

    # Tiêu đề chính
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(12)
    title_run = title_para.add_run(title)
    title_run.font.name = "DejaVuSans"
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    _add_horizontal_rule(title_para)

    lines = content_md.splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            lvl = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            heading_para = doc.add_paragraph()
            heading_para.paragraph_format.space_before = Pt(10)
            heading_para.paragraph_format.space_after = Pt(4)
            run = heading_para.add_run(text)
            run.font.name = "DejaVuSans"
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            run.font.size = Pt(18 - (lvl - 1) * 2)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            hr_para = doc.add_paragraph()
            _add_horizontal_rule(hr_para)
            i += 1
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i]):
                text = re.sub(r"^\s*[-*+]\s+", "", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, text, base_size=11, base_font="DejaVuSans")
                i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                text = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                p = doc.add_paragraph(style="List Number")
                _add_inline(p, text, base_size=11, base_font="DejaVuSans")
                i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            text = re.sub(r"^>\s?", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(text)
            run.font.name = "DejaVuSans"
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            continue

        # Empty line
        if not line.strip():
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_inline(p, line, base_size=11, base_font="DejaVuSans")
        i += 1

    if in_code and code_buffer:
        _add_code_block(doc, "\n".join(code_buffer))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers DOCX ──────────────────────────────────────────────

def _set_doc_default_font(doc, font_name: str, font_size_pt: int):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_elem.insert(0, doc_defaults)

    rPrDefault = doc_defaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        doc_defaults.append(rPrDefault)

    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), str(font_size_pt * 2))


def _add_code_block(doc, code_text: str):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)

    run = p.add_run(code_text)
    run.font.name = "DejaVuSansMono"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def _add_horizontal_rule(para):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F46E5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_inline(paragraph, text: str, base_size: int = 11, base_font: str = "DejaVuSans"):
    from docx.shared import Pt, RGBColor

    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.name = base_font
            r.font.size = Pt(base_size)
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
            r.font.name = base_font
            r.font.size = Pt(base_size)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "DejaVuSansMono"
            r.font.size = Pt(9)
        else:
            r = paragraph.add_run(part)
            r.font.name = base_font
            r.font.size = Pt(base_size)


# ═════════════════════════════════════════════════════════════
# DISPATCHER
# ═════════════════════════════════════════════════════════════
def export_document(title: str, content_md: str, fmt: str) -> tuple[bytes, str, str]:
    safe_title = re.sub(r"[^\w\-_. ]", "_", title or "document").strip() or "document"
    fmt = (fmt or "").lower()

    if fmt == "md":
        return (export_markdown(content_md),
                "text/markdown; charset=utf-8",
                f"{safe_title}.md")

    if fmt == "pdf":
        return (export_pdf(title, content_md),
                "application/pdf",
                f"{safe_title}.pdf")

    if fmt == "docx":
        return (export_docx(title, content_md),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{safe_title}.docx")

    raise ValueError(f"Định dạng không hỗ trợ: {fmt}. Chỉ chấp nhận md, pdf, docx.")